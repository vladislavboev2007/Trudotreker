from database import *
from sqlalchemy import func
from sqlalchemy.orm import Session
from fastapi import Depends, FastAPI, Body, HTTPException, UploadFile, BackgroundTasks, File, Form
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
import datetime as dt
from typing import Optional
from pydantic import BaseModel
from docx import Document
import uvicorn
import os
import io
import uuid
from datetime import datetime
import pandas as pd
from openpyxl import Workbook, load_workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.drawing.image import Image
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.charts.piecharts import Pie
from reportlab.graphics.charts.legends import Legend
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import matplotlib.pyplot as plt

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Создаем буферную папку
BUFFER_DIR = "buffer"
os.makedirs(BUFFER_DIR, exist_ok=True)

# Модели для запросов
class TaskCreate(BaseModel):
    name: str
    description: Optional[str] = None
    empId: int


class EmployeeCreate(BaseModel):
    FIO: str


class TaskUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None

# Добавьте эти модели Pydantic для импорта
class ImportResult(BaseModel):
    total_rows: int
    imported_rows: int
    skipped_rows: int
    errors: list[str]

app = FastAPI()

# Добавьте CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Монтируем статические файлы
app.mount("/public", StaticFiles(directory=os.path.join(BASE_DIR, "public")), name="public")
app.mount("/static", StaticFiles(directory=os.path.join(BASE_DIR, "static")), name="static")

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# Маршруты для страниц
@app.get("/")
def main():
    return FileResponse("public/main.html")


@app.get("/addtask")
def add_task_page():
    return FileResponse("public/addtask.html")


@app.get("/addemployee")
def add_employee_page():
    return FileResponse("public/addemployee.html")


@app.get("/employees")
def employees_page():
    return FileResponse("public/employeeslist.html")


@app.get("/statistic")
def statistic_page():
    return FileResponse("public/statistic.html")

@app.get("/employeeslist.html") # <-- Добавьте эту строку
def employees_list_page():
    return FileResponse("public/employeeslist.html")

# API endpoints
@app.get("/api/tasks")
def get_tasks(db: Session = Depends(get_db)):
    tasks = db.query(Task).join(Employee).all()
    result = []
    for task in tasks:
        # Форматируем время в HH:MM:SS
        start_time = task.start.strftime("%H:%M:%S") if task.start else None
        final_time = task.final.strftime("%H:%M:%S") if task.final else None

        # Вычисляем общее время
        total_time = "00:00:00"
        if task.start and task.final:
            start_dt = dt.datetime.combine(dt.date.today(), task.start)
            final_dt = dt.datetime.combine(dt.date.today(), task.final)
            duration = final_dt - start_dt
            total_seconds = int(duration.total_seconds())
            hours = total_seconds // 3600
            minutes = (total_seconds % 3600) // 60
            seconds = total_seconds % 60
            total_time = f"{hours:02d}:{minutes:02d}:{seconds:02d}"

        result.append({
            "taskId": task.taskId,
            "name": task.name,
            "description": task.description,
            "start": start_time,
            "final": final_time,
            "total_time": total_time,
            "date": str(task.date) if task.date else None,
            "FIO": task.employee.FIO
        })
    return result


@app.get("/api/employees")
def get_employees(db: Session = Depends(get_db)):
    employees = db.query(Employee).all()
    return [{"empId": emp.empId, "FIO": emp.FIO} for emp in employees]


@app.post("/api/tasks")
def create_task(task: TaskCreate, db: Session = Depends(get_db)):
    # Проверяем существование сотрудника
    employee = db.query(Employee).filter(Employee.empId == task.empId).first()
    if not employee:
        raise HTTPException(status_code=404, detail="Сотрудник не найден")

    db_task = Task(
        name=task.name,
        description=task.description,
        empId=task.empId,
        date=dt.date.today()
    )
    db.add(db_task)
    db.commit()
    db.refresh(db_task)
    return {"message": "Задача создана", "taskId": db_task.taskId}


@app.post("/api/employees")
def create_employee(employee: EmployeeCreate, db: Session = Depends(get_db)):
    try:
        db_employee = Employee(FIO=employee.FIO) # Значение empId не передается, генерируется БД
        db.add(db_employee)
        db.commit()
        db.refresh(db_employee) # Обновляем объект, чтобы получить сгенерированный empId
        return {"message": "Сотрудник добавлен", "empId": db_employee.empId}
    except Exception as e:
        db.rollback() # Важно: откатываем транзакцию в случае ошибки
        print(f"Ошибка при добавлении сотрудника в БД: {e}") # Сообщение появится в логах сервера
        raise HTTPException(status_code=500, detail="Внутренняя ошибка сервера")


@app.put("/api/tasks/{task_id}/start")
def start_task(task_id: int, db: Session = Depends(get_db)):
    task = db.query(Task).filter(Task.taskId == task_id).first()
    if not task:
        raise HTTPException(status_code=404, detail="Задача не найдена")

    task.start = dt.datetime.now().time()
    db.commit()
    return {"message": "Задача начата"}


@app.put("/api/tasks/{task_id}/stop")
def stop_task(task_id: int, db: Session = Depends(get_db)):
    task = db.query(Task).filter(Task.taskId == task_id).first()
    if not task:
        raise HTTPException(status_code=404, detail="Задача не найдена")

    task.final = dt.datetime.now().time()
    db.commit()
    return {"message": "Задача остановлена"}


@app.delete("/api/tasks/{task_id}")
def delete_task(task_id: int, db: Session = Depends(get_db)):
    task = db.query(Task).filter(Task.taskId == task_id).first()
    if not task:
        raise HTTPException(status_code=404, detail="Задача не найдена")

    db.delete(task)
    db.commit()
    return {"message": "Задача удалена"}


@app.delete("/api/employees/{emp_id}")
def delete_employee(emp_id: int, db: Session = Depends(get_db)):
    employee = db.query(Employee).filter(Employee.empId == emp_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="Сотрудник не найден")

    # Проверяем, есть ли задачи у сотрудника
    tasks = db.query(Task).filter(Task.empId == emp_id).all()
    if tasks:
        raise HTTPException(status_code=400, detail="Нельзя удалить сотрудника с задачами")

    db.delete(employee)
    db.commit()
    return {"message": "Сотрудник удален"}


@app.get("/api/statistics")
def get_statistics(db: Session = Depends(get_db)):
    # Общее время работы
    tasks = db.query(Task).filter(Task.final.isnot(None)).all()
    total_seconds = 0

    for task in tasks:
        if task.start and task.final:
            start_dt = dt.datetime.combine(dt.date.today(), task.start)
            final_dt = dt.datetime.combine(dt.date.today(), task.final)
            total_seconds += (final_dt - start_dt).total_seconds()

    total_hours = int(total_seconds // 3600)
    total_minutes = int((total_seconds % 3600) // 60)
    total_seconds_remaining = int(total_seconds % 60)
    total_time = f"{total_hours:02d}:{total_minutes:02d}:{total_seconds_remaining:02d}"

    # Статистика по задачам
    total_tasks = db.query(Task).count()
    completed_tasks = db.query(Task).filter(Task.final.isnot(None)).count()

    # Работы по дням (последние 7 дней)
    seven_days_ago = dt.date.today() - dt.timedelta(days=7)
    recent_tasks = db.query(Task).join(Employee).filter(
        Task.date >= seven_days_ago,
        Task.final.isnot(None)
    ).all()

    daily_work = []
    for task in recent_tasks:
        if task.start and task.final:
            start_dt = dt.datetime.combine(dt.date.today(), task.start)
            final_dt = dt.datetime.combine(dt.date.today(), task.final)
            duration_seconds = (final_dt - start_dt).total_seconds()
            hours = int(duration_seconds // 3600)
            minutes = int((duration_seconds % 3600) // 60)
            seconds = int(duration_seconds % 60)

            daily_work.append({
                "time": f"{hours:02d}:{minutes:02d}:{seconds:02d}",
                "date": str(task.date),
                "FIO": task.employee.FIO
            })

    return {
        "total_time": total_time,
        "total_tasks": total_tasks,
        "completed_tasks": completed_tasks,
        "daily_work": daily_work
    }


def delete_file(filepath: str):
    """Удаление файла после отправки"""
    try:
        if os.path.exists(filepath):
            os.remove(filepath)
    except Exception as e:
        print(f"Ошибка удаления файла {filepath}: {e}")


from docx.oxml.ns import qn
from docx.shared import Pt



def get_family_with_initials(full_name):
    """
    Преобразует полное ФИО в фамилию и инициалы
    Пример: 'Колпаков Матвей Николаевич' -> 'Колпаков М.Н.'
    """
    if not full_name or not isinstance(full_name, str):
        return full_name

    parts = full_name.strip().split()

    if len(parts) < 2:
        return full_name

    # Фамилия - первое слово
    family = parts[0]

    # Инициалы из остальных слов
    initials = ''.join([f"{part[0]}." for part in parts[1:]])

    return f"{family}.{initials}"


def replace_in_paragraph(paragraph, old, new, bold=False):
    """Замена текста с сохранением форматирования и возможностью установки жирного шрифта"""
    if old in paragraph.text:
        # Сохраняем все runs и их форматирование
        runs_to_keep = []
        for run in paragraph.runs:
            if old not in run.text:
                runs_to_keep.append(run)
            else:
                # Заменяем текст в run
                new_run = paragraph.add_run(run.text.replace(old, str(new)))
                # Копируем форматирование из оригинального run
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.bold = bold if bold else run.font.bold
                new_run.font.italic = run.font.italic
                new_run.font.underline = run.font.underline
                # Для поддержки кириллицы
                new_run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)

        # Удаляем старые runs
        for run in paragraph.runs[:]:
            run_element = run._element
            run_element.getparent().remove(run_element)

        # Добавляем сохраненные runs обратно
        for run in runs_to_keep:
            new_run = paragraph.add_run(run.text)
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.bold = run.font.bold
            new_run.font.italic = run.font.italic
            new_run.font.underline = run.font.underline
            new_run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)

@app.get("/api/employees/{emp_id}/report")
def generate_employee_report(
        emp_id: int,
        background_tasks: BackgroundTasks,
        db: Session = Depends(get_db)
):
    try:
        employee = db.query(Employee).filter(Employee.empId == emp_id).first()
        if not employee:
            raise HTTPException(status_code=404, detail="Сотрудник не найден")

        tasks = db.query(Task).filter(
            Task.empId == emp_id,
            Task.start.isnot(None),
            Task.final.isnot(None)
        ).order_by(Task.date).all()

        if not tasks:
            raise HTTPException(status_code=400, detail="Нет завершенных задач для сотрудника")

        # Собираем данные
        task_list = []
        total_seconds = 0
        dates = []
        for task in tasks:
            start_dt = dt.datetime.combine(dt.date.today(), task.start)
            final_dt = dt.datetime.combine(dt.date.today(), task.final)
            duration = final_dt - start_dt
            hours = duration.total_seconds() / 3600
            total_seconds += duration.total_seconds()
            if task.date:
                dates.append(task.date)
            task_list.append({
                "id": task.taskId,
                "name": task.name,
                "description": task.description or "",
                "hours": f"{hours:.2f}"
            })

        total_hours = total_seconds / 3600
        start_date = min(dates).strftime('%d.%m.%Y') if dates else 'Н/Д'
        end_date = max(dates).strftime('%d.%m.%Y') if dates else 'Н/Д'

        # Получаем фамилию с инициалами
        family_with_inits = get_family_with_initials(employee.FIO)

        # Загружаем шаблон
        template_path = "report_template.docx"
        if not os.path.exists(template_path):
            raise HTTPException(status_code=500, detail="Шаблон report_template.docx не найден")

        doc = Document(template_path)

        # === Замена простых переменных ===
        def replace_in_paragraph(paragraph, old, new):
            """Замена текста с явным указанием размера шрифта"""
            if old in paragraph.text:
                original_text = paragraph.text
                paragraph.text = paragraph.text.replace(old, str(new))

                # Устанавливаем размер шрифта для всех runs в параграфе
                for run in paragraph.runs:
                    run.font.size = Pt(14)
                    run.font.name = "Times New Roman"
                    # Для поддержки кириллицы
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, "{{ fio }}", employee.FIO)
            replace_in_paragraph(paragraph, "{{ start_date }}", start_date)
            replace_in_paragraph(paragraph, "{{ end_date }}", end_date)
            replace_in_paragraph(paragraph, "{{ total_hours }}", f"{total_hours:.2f}")
            replace_in_paragraph(paragraph, "{{ family_with_inits }}", family_with_inits)

        # === Работа с таблицей ===
        if len(doc.tables) == 0:
            raise HTTPException(status_code=500, detail="В шаблоне не найдена таблица")

        table = doc.tables[0]

        # Удаляем первую строку данных (если она есть и не является заголовком)
        if len(table.rows) > 1:
            row = table.rows[1]._element
            row.getparent().remove(row)

        # Функция для установки шрифта в ячейке таблицы
        def set_cell_font(cell, font_name="Times New Roman", font_size=14):
            """Устанавливает шрифт для всех параграфов в ячейке таблицы"""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        # Добавляем строки с данными и устанавливаем шрифт
        for task in task_list:
            row_cells = table.add_row().cells
            row_cells[0].text = str(task["id"])
            row_cells[1].text = task["name"]
            row_cells[2].text = task["description"]
            row_cells[3].text = task["hours"]

            # Устанавливаем шрифт для всех ячеек новой строки
            for cell in row_cells:
                set_cell_font(cell, "Times New Roman", 14)

        # Дополнительно: устанавливаем шрифт для заголовков таблицы (если нужно)
        if len(table.rows) > 0:
            header_cells = table.rows[0].cells
            for cell in header_cells:
                set_cell_font(cell, "Times New Roman", 14)
                # Можно сделать заголовки жирными
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        # Сохраняем
        filename = f"report_{uuid.uuid4().hex}.docx"
        filepath = os.path.join(BUFFER_DIR, filename)
        doc.save(filepath)

        background_tasks.add_task(delete_file, filepath)
        download_filename = f"report_{employee.FIO.replace(' ', '_')}.docx"
        return FileResponse(
            path=filepath,
            filename=download_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Ошибка генерации отчёта: {str(e)}")


@app.get("/api/employees/{emp_id}/excel-report")
def generate_employee_excel_report(
        emp_id: int,
        background_tasks: BackgroundTasks,
        db: Session = Depends(get_db)
):
    try:
        employee = db.query(Employee).filter(Employee.empId == emp_id).first()
        if not employee:
            raise HTTPException(status_code=404, detail="Сотрудник не найден")

        # Получаем ВСЕ задачи сотрудника (не только завершенные)
        all_tasks = db.query(Task).filter(Task.empId == emp_id).order_by(Task.date).all()

        if not all_tasks:
            raise HTTPException(status_code=400, detail="Нет задач для сотрудника")

        # Собираем данные для таблицы задач
        task_data = []
        total_seconds = 0
        completed_tasks_count = 0
        dates = []

        for task in all_tasks:
            hours = 0
            if task.start and task.final:
                if task.date:
                    start_dt = dt.datetime.combine(task.date, task.start)
                    final_dt = dt.datetime.combine(task.date, task.final)
                else:
                    start_dt = dt.datetime.combine(dt.date.today(), task.start)
                    final_dt = dt.datetime.combine(dt.date.today(), task.final)

                duration = final_dt - start_dt
                hours = duration.total_seconds() / 3600
                total_seconds += duration.total_seconds()
                completed_tasks_count += 1

            if task.date:
                dates.append(task.date)

            task_data.append({
                "№": task.taskId,
                "Название задания": task.name,
                "Описание задачи": task.description or "",
                "Дата выполнения": task.date.strftime('%d.%m.%Y') if task.date else '',
                "Время выполнения (в час)": round(hours, 2)
            })

        # Рассчитываем показатели
        total_hours = total_seconds / 3600
        total_tasks_count = len(all_tasks)
        completion_rate = (completed_tasks_count / total_tasks_count * 100) if total_tasks_count > 0 else 0

        start_date = min(dates).strftime('%d.%m.%Y') if dates else 'Н/Д'
        end_date = max(dates).strftime('%d.%m.%Y') if dates else 'Н/Д'

        # Получаем фамилию с инициалами
        family_with_inits = get_family_with_initials(employee.FIO)

        # Загружаем шаблон из static папки
        template_path = "report_template.xlsx"
        if not os.path.exists(template_path):
            raise HTTPException(status_code=500, detail="Шаблон static/report_template.xlsx не найден")

        # Открываем шаблон
        wb = load_workbook(template_path)
        ws = wb.active

        # Заполняем основные данные согласно шаблону
        ws['C9'] = employee.FIO  # [FIO] - ФИО сотрудника
        ws['C10'] = f"{start_date} - {end_date}"  # [min(date)-max(date)] - Период
        ws['C11'] = round(total_hours, 2)  # [total_hours] - Общее кол-во часов
        ws['C12'] = total_tasks_count  # [total count tasks] - Общее кол-во задач
        ws['C13'] = f"{completion_rate:.1f}%"  # [percent of completed tasks] - Процент выполнения
        ws['C14'] = completed_tasks_count  # [total count completed tasks] - Выполнено задач

        # Выравниваем по левому краю ячейки A5:B5 (адрес)
        ws['A5'].alignment = Alignment(horizontal='left')
        ws['B5'].alignment = Alignment(horizontal='left')

        # Выравниваем по левому краю числовые ячейки C11, C12, C14
        ws['C11'].alignment = Alignment(horizontal='left')
        ws['C12'].alignment = Alignment(horizontal='left')
        ws['C14'].alignment = Alignment(horizontal='left')

        # Заполняем таблицу задач начиная с 17 строки
        start_row = 17

        # Устанавливаем заголовки таблицы (строка 17)
        headers = ["№", "Название задания", "Описание задачи", "Дата выполнения", "Время выполнения (в час)"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=start_row, column=col, value=header)
            # Стили для заголовков
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # Заполняем данные задач (начиная со строки 18)
        for i, task in enumerate(task_data):
            row = start_row + 1 + i

            # Заполняем ячейки
            ws.cell(row=row, column=1, value=task["№"])
            ws.cell(row=row, column=2, value=task["Название задания"])
            ws.cell(row=row, column=3, value=task["Описание задачи"])
            ws.cell(row=row, column=4, value=task["Дата выполнения"])
            ws.cell(row=row, column=5, value=task["Время выполнения (в час)"])

            # Добавляем границы для всех ячеек
            for col in range(1, 6):
                cell = ws.cell(row=row, column=col)
                cell.border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
                # Выравнивание по центру для №, даты и времени
                if col in [1, 4, 5]:
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                else:
                    cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)

        # Добавляем подпись после таблицы задач
        last_task_row = start_row + len(task_data) + 2

        # Подпись ответственного сотрудника (выравниваем по правому краю в столбце E)
        signature_cell = ws.cell(row=last_task_row, column=5, value="Подпись ответственного сотрудника")
        signature_cell.alignment = Alignment(horizontal='right')

        # ФИО с инициалами (выравниваем по правому краю в столбце E)
        name_cell = ws.cell(row=last_task_row + 1, column=5, value=f"__________ / {family_with_inits}")
        name_cell.alignment = Alignment(horizontal='right')

        # Настраиваем ширину колонок для лучшего отображения
        column_widths = [8, 30, 40, 17, 24]  # Обновленные ширины для столбцов D и E
        for i, width in enumerate(column_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = width

        # Сохраняем файл
        filename = f"excel_report_{uuid.uuid4().hex}.xlsx"
        filepath = os.path.join(BUFFER_DIR, filename)
        wb.save(filepath)

        background_tasks.add_task(delete_file, filepath)
        download_filename = f"excel_report_{employee.FIO.replace(' ', '_')}.xlsx"

        return FileResponse(
            path=filepath,
            filename=download_filename,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.document"
        )

    except Exception as e:
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Ошибка генерации Excel отчёта: {str(e)}")


# Регистрируем шрифты для поддержки кириллицы (добавьте в начало файла)
def register_fonts():
    """Регистрирует шрифты для поддержки кириллицы"""
    try:
        # Попробуем найти стандартные шрифты с поддержкой кириллицы
        font_paths = [
            # Windows
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/times.ttf",
            # Linux
            "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
            # MacOS
            "/Library/Fonts/Arial.ttf",
            # Текущая директория
            "arial.ttf",
            "times.ttf"
        ]

        # Регистрируем Arial
        for font_path in font_paths:
            if os.path.exists(font_path) and "arial" in font_path.lower():
                pdfmetrics.registerFont(TTFont('Arial', font_path))
                pdfmetrics.registerFont(TTFont('Arial-Bold', font_path))
                print(f"Зарегистрирован шрифт Arial: {font_path}")
                break
        else:
            # Если Arial не найден, используем стандартный шрифт
            print("Шрифт Arial не найден, используется стандартный")

        # Регистрируем Times New Roman
        for font_path in font_paths:
            if os.path.exists(font_path) and "times" in font_path.lower():
                pdfmetrics.registerFont(TTFont('Times-Roman', font_path))
                pdfmetrics.registerFont(TTFont('Times-Bold', font_path))
                print(f"Зарегистрирован шрифт Times: {font_path}")
                break

    except Exception as e:
        print(f"Ошибка регистрации шрифтов: {e}")


# Вызов регистрации шрифтов при запуске
register_fonts()


@app.get("/api/employees/{emp_id}/pdf-report")
def generate_employee_pdf_report(
        emp_id: int,
        background_tasks: BackgroundTasks,
        db: Session = Depends(get_db)
):
    try:
        employee = db.query(Employee).filter(Employee.empId == emp_id).first()
        if not employee:
            raise HTTPException(status_code=404, detail="Сотрудник не найден")

        # Получаем задачи за последние 30 дней
        thirty_days_ago = dt.date.today() - dt.timedelta(days=30)
        tasks = db.query(Task).filter(
            Task.empId == emp_id,
            Task.start.isnot(None),
            Task.final.isnot(None),
            Task.date >= thirty_days_ago
        ).order_by(Task.date).all()

        if not tasks:
            raise HTTPException(status_code=400, detail="Нет завершенных задач для сотрудника за последний месяц")

        # Собираем данные
        task_data = []
        total_seconds = 0
        dates = []

        # Группируем задачи по датам для гистограммы
        date_hours = {}

        for task in tasks:
            if task.start and task.final:
                if task.date:
                    start_dt = dt.datetime.combine(task.date, task.start)
                    final_dt = dt.datetime.combine(task.date, task.final)
                else:
                    start_dt = dt.datetime.combine(dt.date.today(), task.start)
                    final_dt = dt.datetime.combine(dt.date.today(), task.final)

                duration = final_dt - start_dt
                hours = duration.total_seconds() / 3600
                total_seconds += duration.total_seconds()

                if task.date:
                    dates.append(task.date)
                    date_str = task.date.strftime('%d.%m.%Y')
                    # Суммируем часы по датам
                    date_hours[date_str] = date_hours.get(date_str, 0) + hours

                task_data.append({
                    "№": task.taskId,
                    "Название задания": task.name,
                    "Описание": task.description or "",
                    "Дата выполнения": task.date.strftime('%d.%m.%Y') if task.date else '',
                    "Время выполнения (в час)": round(hours, 2)
                })

        total_hours = total_seconds / 3600
        start_date = min(dates).strftime('%d.%m.%Y') if dates else 'Н/Д'
        end_date = max(dates).strftime('%d.%m.%Y') if dates else 'Н/Д'

        # Рассчитываем продолжительность периода в днях
        count_days = 0
        if dates:
            start_dt = min(dates)
            end_dt = max(dates)
            count_days = (end_dt - start_dt).days + 1  # +1 чтобы включить начальный день

        # Функция для правильного склонения слова "день"
        def get_day_form(days):
            if days % 10 == 1 and days % 100 != 11:
                return "день"
            elif 2 <= days % 10 <= 4 and (days % 100 < 10 or days % 100 >= 20):
                return "дня"
            else:
                return "дней"

        # Формируем строку периода с продолжительностью
        if dates:
            day_form = get_day_form(count_days)
            period_text = f"{start_date} - {end_date} ({count_days} {day_form})"
        else:
            period_text = "Н/Д"

        family_with_inits = get_family_with_initials(employee.FIO)

        # Создаем PDF файл
        filename = f"pdf_report_{uuid.uuid4().hex}.pdf"
        filepath = os.path.join(BUFFER_DIR, filename)

        # Регистрируем шрифты Times New Roman
        try:
            # Пути к файлам шрифтов Times New Roman
            times_font_paths = [
                "static/ttf/times.ttf",
                "static/ttf/times-new-roman.ttf",
                "static/ttf/times-roman.ttf",
                "C:/Windows/Fonts/times.ttf",
                "C:/Windows/Fonts/timesbd.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf"
            ]

            times_found = False
            times_bold_found = False

            for font_path in times_font_paths:
                if os.path.exists(font_path):
                    if "bd" in font_path.lower() or "bold" in font_path.lower():
                        pdfmetrics.registerFont(TTFont('Times-Bold', font_path))
                        times_bold_found = True
                    else:
                        pdfmetrics.registerFont(TTFont('Times-Roman', font_path))
                        times_found = True

            # Если не нашли шрифты, используем DejaVuSans
            if not times_found:
                if os.path.exists("static/ttf/DejaVoSans.ttf"):
                    pdfmetrics.registerFont(TTFont('Times-Roman', "static/ttf/DejaVoSans.ttf"))
                    times_found = True

            if not times_bold_found:
                if os.path.exists("static/ttf/DejaVoSans-Bold.ttf"):
                    pdfmetrics.registerFont(TTFont('Times-Bold', "static/ttf/DejaVoSans-Bold.ttf"))
                    times_bold_found = True

        except Exception as e:
            print(f"Ошибка регистрации шрифтов: {e}")

        # Создаем документ
        doc = SimpleDocTemplate(filepath, pagesize=A4,
                                leftMargin=2 * cm, rightMargin=2 * cm,
                                topMargin=2 * cm, bottomMargin=2 * cm)
        elements = []
        styles = getSampleStyleSheet()

        # Создаем кастомные стили с Times New Roman размером 14 и межстрочным интервалом 1.15
        line_height = 14 * 1.15  # 1.15 межстрочный интервал

        try:
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName='Times-Bold',
                fontSize=14,
                leading=line_height,
                spaceAfter=20,
                alignment=1,  # CENTER
                textColor=colors.black
            )
            header_left_style = ParagraphStyle(
                'HeaderLeftStyle',
                parent=styles['Normal'],
                fontName='Times-Roman',
                fontSize=14,
                leading=line_height,
                spaceAfter=30,
                alignment=0  # LEFT
            )
            header_right_style = ParagraphStyle(
                'HeaderRightStyle',
                parent=styles['Normal'],
                fontName='Times-Roman',
                fontSize=14,
                leading=line_height,
                spaceAfter=30,
                alignment=2  # RIGHT
            )
            header_bold_style = ParagraphStyle(
                'HeaderBoldStyle',
                parent=styles['Normal'],
                fontName='Times-Bold',
                fontSize=14,
                leading=line_height,
                spaceAfter=6
            )
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName='Times-Roman',
                fontSize=14,
                leading=line_height,
                spaceAfter=6
            )
            table_header_style = ParagraphStyle(
                'TableHeader',
                parent=styles['Normal'],
                fontName='Times-Bold',
                fontSize=14,
                leading=line_height,
                alignment=1  # CENTER
            )
            table_data_style = ParagraphStyle(
                'TableData',
                parent=styles['Normal'],
                fontName='Times-Roman',
                fontSize=14,
                leading=line_height
            )
            signature_style = ParagraphStyle(
                'SignatureStyle',
                parent=styles['Normal'],
                fontName='Times-Roman',
                fontSize=14,
                leading=line_height,
                alignment=2,  # RIGHT
                spaceBefore=20
            )
        except:
            # Резервные стили если Times не доступен
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=14,
                leading=line_height,
                spaceAfter=20,
                alignment=1,
                textColor=colors.black
            )
            header_left_style = ParagraphStyle(
                'HeaderLeftStyle',
                parent=styles['Normal'],
                fontSize=14,
                leading=line_height,
                spaceAfter=20,
                alignment=0
            )
            header_right_style = ParagraphStyle(
                'HeaderRightStyle',
                parent=styles['Normal'],
                fontSize=14,
                leading=line_height,
                spaceAfter=20,
                alignment=2
            )
            header_bold_style = ParagraphStyle(
                'HeaderBoldStyle',
                parent=styles['Normal'],
                fontSize=14,
                leading=line_height,
                spaceAfter=6
            )
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=14,
                leading=line_height,
                spaceAfter=6
            )
            table_header_style = ParagraphStyle(
                'TableHeader',
                parent=styles['Normal'],
                fontSize=14,
                leading=line_height,
                alignment=1
            )
            table_data_style = ParagraphStyle(
                'TableData',
                parent=styles['Normal'],
                fontSize=14,
                leading=line_height
            )
            signature_style = ParagraphStyle(
                'SignatureStyle',
                parent=styles['Normal'],
                fontSize=14,
                leading=line_height,
                alignment=2,
                spaceBefore=20
            )

        # === ВЕРХНИЕ ШАПКИ ===
        header_data = [
            [
                # Левая шапка - выровнена по левому краю
                Paragraph('ООО "1C:Максималист"<br/>Адрес:<br/>г.Москва, проспект Мира, д.126<br/>ИНН 2332543567',
                          header_left_style),
                # Правая шапка - выровнена по правому краю
                Paragraph(
                    'УТВЕРЖДАЮ<br/>Руководитель подразделения<br/>_____________ / Шувалов.С.Н<br/>от «__» ___________ 20__ г.',
                    header_right_style)
            ]
        ]

        header_table = Table(header_data, colWidths=[10 * cm, 10 * cm])
        header_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elements.append(header_table)
        elements.append(Spacer(1, 1 * cm))

        # === ЗАГОЛОВОК ===
        title = Paragraph("ОТЧЕТ О ВЫПОЛНЕННЫХ ЗАДАЧАХ СОТРУДНИКА", title_style)
        elements.append(title)
        elements.append(Spacer(1, 1 * cm))

        # === ИНФОРМАЦИЯ О СОТРУДНИКЕ ===
        employee_info_data = [
            [Paragraph('ФИО сотрудника:', header_bold_style), Paragraph(employee.FIO, normal_style)],
            [Paragraph('Период анализа:', header_bold_style), Paragraph(period_text, normal_style)],
            [Paragraph('Общее время работы:', header_bold_style), Paragraph(f"{total_hours:.2f} часов", normal_style)],
            [Paragraph('Количество задач:', header_bold_style), Paragraph(str(len(tasks)), normal_style)]
        ]

        employee_table = Table(employee_info_data, colWidths=[5 * cm, 10 * cm])
        employee_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elements.append(employee_table)
        elements.append(Spacer(1, 1 * cm))

        # === ГИСТОГРАММА ===
        if date_hours:
            elements.append(Paragraph("График выполнения задач по датам", header_bold_style))
            elements.append(Spacer(1, 0.5 * cm))

            # Правильно сортируем даты по хронологии
            def parse_date(date_str):
                return dt.datetime.strptime(date_str, '%d.%m.%Y')

            # Сортируем даты по возрастанию (от старых к новым)
            dates_sorted = sorted(date_hours.keys(), key=parse_date)
            hours_data = [date_hours[date] for date in dates_sorted]

            # Создаем рисунок для гистограммы
            drawing = Drawing(400, 200)
            chart = VerticalBarChart()
            chart.x = 30
            chart.y = 50
            chart.height = 150
            chart.width = 420
            chart.data = [hours_data]
            chart.categoryAxis.categoryNames = dates_sorted
            chart.categoryAxis.labels.angle = 45
            chart.categoryAxis.labels.dy = -12
            chart.categoryAxis.labels.fontName = 'Times-Roman'
            chart.categoryAxis.labels.fontSize = 8
            chart.valueAxis.valueMin = 0
            chart.valueAxis.valueMax = max(hours_data) * 1.2 if hours_data else 10
            chart.valueAxis.labels.fontName = 'Times-Roman'
            chart.valueAxis.labels.fontSize = 8
            chart.bars[0].fillColor = colors.HexColor('#FF0000')
            chart.barLabels.fontName = 'Times-Roman'
            chart.barLabels.fontSize = 8

            drawing.add(chart)
            elements.append(drawing)
            elements.append(Spacer(1, 0.5 * cm))

        # === ТАБЛИЦА ЗАДАЧ ===
        elements.append(Paragraph("Список выполненных задач", header_bold_style))
        elements.append(Spacer(1, 0.5 * cm))

        # Заголовки таблицы
        task_table_data = [[
            Paragraph('№', table_header_style),
            Paragraph('Название задания', table_header_style),
            Paragraph('Описание', table_header_style),
            Paragraph('Дата выполнения', table_header_style),
            Paragraph('Время (часы)', table_header_style)
        ]]

        # Данные таблицы
        for task in task_data:
            task_table_data.append([
                Paragraph(str(task["№"]), table_data_style),
                Paragraph(task["Название задания"], table_data_style),
                Paragraph(task["Описание"], table_data_style),
                Paragraph(task["Дата выполнения"], table_data_style),
                Paragraph(str(task["Время выполнения (в час)"]), table_data_style)
            ])

        task_table = Table(task_table_data, colWidths=[1.5 * cm, 5 * cm, 5 * cm, 3 * cm, 2.5 * cm])
        task_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#F0F0F0')),  # Светло-серый фон заголовков
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),  # Черный текст для заголовков
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
            ('ALIGN', (4, 1), (4, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Times-Bold'),
            ('FONTNAME', (0, 1), (-1, -1), 'Times-Roman'),
            ('FONTSIZE', (0, 0), (-1, -1), 14),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#DEE2E6')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        elements.append(task_table)

        # === ПОДПИСЬ ===
        elements.append(Spacer(1, 1 * cm))
        signature = Paragraph(f"Подпись ответственного сотрудника<br/>_____________ / {family_with_inits}",
                              signature_style)
        elements.append(signature)

        # Строим PDF
        doc.build(elements)

        background_tasks.add_task(delete_file, filepath)
        download_filename = f"pdf_report_{employee.FIO.replace(' ', '_')}.pdf"

        return FileResponse(
            path=filepath,
            filename=download_filename,
            media_type="application/pdf"
        )

    except Exception as e:
        import traceback
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Ошибка генерации PDF отчёта: {str(e)}")


@app.get("/api/export/tasks")
def export_tasks(db: Session = Depends(get_db)):
    tasks = db.query(Task).join(Employee).all()

    # Подготавливаем данные для экспорта
    data = []
    for task in tasks:
        # Форматируем время в HH:MM:SS
        start_time = task.start.strftime("%H:%M:%S") if task.start else ""
        final_time = task.final.strftime("%H:%M:%S") if task.final else ""

        # Вычисляем общее время
        total_time = "00:00:00"
        if task.start and task.final:
            start_dt = datetime.combine(datetime.today(), task.start)
            final_dt = datetime.combine(datetime.today(), task.final)
            duration = final_dt - start_dt
            total_seconds = int(duration.total_seconds())
            hours = total_seconds // 3600
            minutes = (total_seconds % 3600) // 60
            seconds = total_seconds % 60
            total_time = f"{hours:02d}:{minutes:02d}:{seconds:02d}"

        data.append({
            "ID": task.taskId,
            "Название": task.name,
            "Описание": task.description or "",
            "Начало": start_time,
            "Конец": final_time,
            "Общее время": total_time,
            "Дата": task.date.strftime("%Y-%m-%d") if task.date else "",
            "ФИО сотрудника": task.employee.FIO
        })

    # Создаем DataFrame
    df = pd.DataFrame(data)

    # Создаем Excel файл в памяти
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Задачи', index=False)

        # Получаем объект листа
        worksheet = writer.sheets['Задачи']

        # Устанавливаем ширину столбцов
        worksheet.column_dimensions['A'].width = 8  # ID
        worksheet.column_dimensions['B'].width = 44  # Название
        worksheet.column_dimensions['C'].width = 50  # Описание
        worksheet.column_dimensions['D'].width = 8  # Начало
        worksheet.column_dimensions['E'].width = 8  # Конец
        worksheet.column_dimensions['F'].width = 14  # Общее время
        worksheet.column_dimensions['G'].width = 10  # Дата
        worksheet.column_dimensions['H'].width = 37  # ФИО сотрудника

        # Включаем перенос по словам для столбцов B и C
        for row in worksheet.iter_rows(min_row=2, max_row=worksheet.max_row, min_col=2, max_col=3):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical='top')

        # Стили для заголовков
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")

        # Применяем стили к заголовкам
        for cell in worksheet[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

            # Добавляем подпись после данных
        last_row = len(df) + 3  # Пропускаем 2 строки после данных

        # Подпись ответственного сотрудника
        signature_cell = worksheet.cell(row=last_row, column=8, value="Подпись руководителя подразделения")
        signature_cell.alignment = Alignment(horizontal='right')

        # ФИО с подписью
        name_cell = worksheet.cell(row=last_row + 1, column=8, value="__________ / Шувалов.С.Н.")
        name_cell.alignment = Alignment(horizontal='right')

    output.seek(0)

    # Возвращаем файл
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=export_tasks.xlsx"}
    )


@app.get("/api/export/employees")
def export_employees(db: Session = Depends(get_db)):
    employees = db.query(Employee).all()

    # Подготавливаем данные для экспорта
    data = []
    for emp in employees:
        data.append({
            "ID": emp.empId,
            "ФИО сотрудника": emp.FIO
        })

    # Создаем DataFrame
    df = pd.DataFrame(data)

    # Создаем Excel файл в памяти
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Сотрудники', index=False)

        # Получаем объект листа
        worksheet = writer.sheets['Сотрудники']

        # Устанавливаем ширину столбцов
        worksheet.column_dimensions['A'].width = 10  # Столбец A (ID) - ширина ~10
        worksheet.column_dimensions['B'].width = 40  # Столбец B (ФИО сотрудника) - ширина ~40

        last_row = len(df) + 3  # Пропускаем 2 строки после данных

        # Подпись ответственного сотрудника
        signature_cell = worksheet.cell(row=last_row, column=2, value="Подпись руководителя подразделения")
        signature_cell.alignment = Alignment(horizontal='right')

        # ФИО с подписью
        name_cell = worksheet.cell(row=last_row + 1, column=2, value="__________ / Шувалов.С.Н.")
        name_cell.alignment = Alignment(horizontal='right')

    output.seek(0)

    # Возвращаем файл
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=export_employees.xlsx"}
    )


@app.post("/api/import/employees")
async def import_employees(
        file: UploadFile = File(...),
        ignore_duplicates: bool = Form(True),
        db: Session = Depends(get_db)
):
    """
    Импорт сотрудников из Excel файла
    """
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Формат файла должен быть .xlsx или .xls")

    try:
        # Читаем файл
        contents = await file.read()
        df = pd.read_excel(io.BytesIO(contents))

        # Нормализуем названия колонок (игнорируем регистр и пробелы)
        df.columns = df.columns.str.strip().str.lower()

        # Ищем колонки с ФИО
        fio_columns = [col for col in df.columns if any(x in col for x in ['fio', 'фио', 'name', 'имя', 'сотрудник'])]
        if not fio_columns:
            raise HTTPException(status_code=400, detail="Не найдена колонка с ФИО сотрудника")

        fio_column = fio_columns[0]
        imported_count = 0
        skipped_count = 0
        errors = []

        # Получаем максимальный существующий ID
        max_id = db.query(func.max(Employee.empId)).scalar() or 0

        for index, row in df.iterrows():
            try:
                fio_value = str(row[fio_column]).strip()
                if not fio_value or pd.isna(fio_value):
                    skipped_count += 1
                    continue

                # Проверяем дубликаты по ФИО
                if ignore_duplicates:
                    existing_employee = db.query(Employee).filter(
                        func.lower(Employee.FIO) == func.lower(fio_value)
                    ).first()
                    if existing_employee:
                        skipped_count += 1
                        continue

                # Создаем нового сотрудника
                new_employee = Employee(FIO=fio_value)
                db.add(new_employee)
                imported_count += 1

            except Exception as e:
                errors.append(f"Строка {index + 2}: {str(e)}")
                skipped_count += 1

        db.commit()

        return ImportResult(
            total_rows=len(df),
            imported_rows=imported_count,
            skipped_rows=skipped_count,
            errors=errors
        )

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Ошибка импорта: {str(e)}")


@app.post("/api/import/tasks")
async def import_tasks(
        file: UploadFile = File(...),
        ignore_duplicates: bool = Form(True),
        db: Session = Depends(get_db)
):
    """
    Импорт задач из Excel файла
    """
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Формат файла должен быть .xlsx или .xls")

    try:
        # Читаем файл
        contents = await file.read()
        df = pd.read_excel(io.BytesIO(contents))

        # Нормализуем названия колонок
        df.columns = df.columns.str.strip().str.lower()

        # Ищем необходимые колонки
        name_columns = [col for col in df.columns if any(x in col for x in ['name', 'название', 'задача'])]
        emp_columns = [col for col in df.columns if
                       any(x in col for x in ['emp', 'employee', 'сотрудник', 'fio', 'фио'])]

        if not name_columns:
            raise HTTPException(status_code=400, detail="Не найдена колонка с названием задачи")
        if not emp_columns:
            raise HTTPException(status_code=400, detail="Не найдена колонка с сотрудником")

        name_column = name_columns[0]
        emp_column = emp_columns[0]

        # Ищем опциональные колонки
        desc_column = next((col for col in df.columns if any(x in col for x in ['desc', 'description', 'описание'])),
                           None)
        start_column = next((col for col in df.columns if any(x in col for x in ['start', 'начало'])), None)
        final_column = next((col for col in df.columns if any(x in col for x in ['final', 'end', 'конец'])), None)
        date_column = next((col for col in df.columns if any(x in col for x in ['date', 'дата'])), None)

        imported_count = 0
        skipped_count = 0
        errors = []

        for index, row in df.iterrows():
            try:
                name_value = str(row[name_column]).strip()
                emp_value = str(row[emp_column]).strip()

                if not name_value or pd.isna(name_value) or not emp_value or pd.isna(emp_value):
                    skipped_count += 1
                    continue

                # Ищем сотрудника по ФИО или ID
                employee = None
                if emp_value.isdigit():
                    employee = db.query(Employee).filter(Employee.empId == int(emp_value)).first()
                else:
                    employee = db.query(Employee).filter(
                        func.lower(Employee.FIO) == func.lower(emp_value)
                    ).first()

                if not employee:
                    errors.append(f"Строка {index + 2}: Сотрудник '{emp_value}' не найден")
                    skipped_count += 1
                    continue

                # Проверяем дубликаты
                if ignore_duplicates:
                    existing_task = db.query(Task).filter(
                        func.lower(Task.name) == func.lower(name_value),
                        Task.empId == employee.empId
                    ).first()
                    if existing_task:
                        skipped_count += 1
                        continue

                # Создаем задачу
                new_task = Task(
                    name=name_value,
                    description=str(row[desc_column]).strip() if desc_column and not pd.isna(
                        row[desc_column]) else None,
                    empId=employee.empId,
                    date=dt.date.today()
                )

                # Обрабатываем время начала
                if start_column and not pd.isna(row[start_column]):
                    if isinstance(row[start_column], str):
                        try:
                            new_task.start = datetime.strptime(row[start_column], '%H:%M:%S').time()
                        except ValueError:
                            try:
                                new_task.start = datetime.strptime(row[start_column], '%H:%M').time()
                            except ValueError:
                                pass
                    elif isinstance(row[start_column], datetime):
                        new_task.start = row[start_column].time()

                # Обрабатываем время окончания
                if final_column and not pd.isna(row[final_column]):
                    if isinstance(row[final_column], str):
                        try:
                            new_task.final = datetime.strptime(row[final_column], '%H:%M:%S').time()
                        except ValueError:
                            try:
                                new_task.final = datetime.strptime(row[final_column], '%H:%M').time()
                            except ValueError:
                                pass
                    elif isinstance(row[final_column], datetime):
                        new_task.final = row[final_column].time()

                # Обрабатываем дату
                if date_column and not pd.isna(row[date_column]):
                    if isinstance(row[date_column], str):
                        try:
                            new_task.date = datetime.strptime(row[date_column], '%Y-%m-%d').date()
                        except ValueError:
                            try:
                                new_task.date = datetime.strptime(row[date_column], '%d.%m.%Y').date()
                            except ValueError:
                                pass
                    elif isinstance(row[date_column], datetime):
                        new_task.date = row[date_column].date()

                db.add(new_task)
                imported_count += 1

            except Exception as e:
                errors.append(f"Строка {index + 2}: {str(e)}")
                skipped_count += 1

        db.commit()

        return ImportResult(
            total_rows=len(df),
            imported_rows=imported_count,
            skipped_rows=skipped_count,
            errors=errors
        )

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Ошибка импорта: {str(e)}")



# Эндпоинты для удаления всех записей
@app.delete("/api/tasks")
def delete_all_tasks(db: Session = Depends(get_db)):
    try:
        db.query(Task).delete()
        db.commit()
        return {"message": "Все задачи удалены"}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Ошибка при удалении: {str(e)}")


@app.delete("/api/employees")
def delete_all_employees(db: Session = Depends(get_db)):
    try:
        # Проверяем, есть ли задачи у сотрудников
        tasks_exist = db.query(Task).first()
        if tasks_exist:
            raise HTTPException(status_code=400, detail="Нельзя удалить сотрудников, у которых есть задачи")

        db.query(Employee).delete()
        db.commit()
        return {"message": "Все сотрудники удалены"}
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Ошибка при удалении: {str(e)}")



if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host="127.0.0.1",
        port=7000,
        reload=True  # ← Добавьте эту строку
    )